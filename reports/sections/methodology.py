from docx import Document

def add_methodology_section(doc: Document):
    doc.add_heading("Methodology", level=1)

    doc.add_paragraph(
        "The penetration test was conducted using a structured and industry-recognized methodology, "
        "primarily based on the OWASP Testing Guide v4. This ensures a thorough evaluation of the target system."
    )

    doc.add_heading("1. Information Gathering", level=2)
    doc.add_paragraph("- DNS Enumeration\n- WHOIS lookups\n- Port Scanning\n- OSINT\n- Banner Grabbing")

    doc.add_heading("2. Threat Modeling", level=2)
    doc.add_paragraph("Identify trust boundaries, data flows, and threat agents.")

    doc.add_heading("3. Vulnerability Analysis", level=2)
    doc.add_paragraph("Manual testing and automated scanning to discover known vulnerabilities.")

    doc.add_heading("4. Exploitation", level=2)
    doc.add_paragraph("Attempting to exploit vulnerabilities to confirm impact and risk.")

    doc.add_heading("5. Post-Exploitation", level=2)
    doc.add_paragraph("Understanding possible lateral movement, privilege escalation, and data compromise.")

    doc.add_heading("6. Reporting", level=2)
    doc.add_paragraph("Clear documentation of findings, risk levels, and remediation recommendations.")

    doc.add_heading("Tools and Standards", level=2)
    doc.add_paragraph(
        "Tools: Nmap, Burp Suite, Nikto, SQLMap, custom scripts\n"
        "Standards: OWASP WSTG, PTES\n"
        "Approach: Black-box / Gray-box / White-box (depending on project scope)"
    )
