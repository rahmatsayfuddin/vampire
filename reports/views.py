import os
import threading
from django.shortcuts import get_object_or_404, redirect
from django.http import HttpResponse, FileResponse
from django.conf import settings
from .models import ReportHistory
from projects.models import Project
from reportlab.pdfgen import canvas  # for PDF
from docx import Document  # for Word
from django.http import FileResponse, Http404

def generate_report(request, project_id, format):
    project = get_object_or_404(Project, id=project_id)
    
    filename = f"{project.project_name}_{format}_{project_id}_{ReportHistory.objects.count() + 1}.{format}"
    report = ReportHistory.objects.create(
        project=project,
        file_name=filename,
        format=format,
        status='loading',
    )

    def generate():
        file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        try:
            if format == 'pdf':
                c = canvas.Canvas(file_path)
                c.drawString(100, 800, f"Project: {project.project_name}")
                c.drawString(100, 780, f"Status: {project.status}")
                c.save()
            elif format == 'docx':
                doc = Document()
                doc.add_heading(f"Project Report: {project.project_name}", 0)
                doc.add_paragraph(f"Status: {project.status}")
                doc.save(file_path)
            report.status = 'done'
        except Exception as e:
            report.status = 'error'
        report.save()

    threading.Thread(target=generate).start()
    return redirect('project_detail', pk=project_id)


def download_report(request, report_id):
    report = get_object_or_404(ReportHistory, id=report_id)
    path = os.path.join(settings.MEDIA_ROOT, 'reports', report.file_name)
    if os.path.exists(path):
        return FileResponse(open(path, 'rb'), as_attachment=True)
    raise Http404("File not found.")

def delete_report(request, report_id):
    report = get_object_or_404(ReportHistory, id=report_id)
    path = os.path.join(settings.MEDIA_ROOT, 'reports', report.file_name)
    if os.path.exists(path):
        os.remove(path)
    report.delete()
    return redirect('project_detail', pk=report.project.id)
