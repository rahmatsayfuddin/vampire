import os
from django.template.loader import render_to_string
from django.conf import settings
from xhtml2pdf import pisa
from docx import Document
from datetime import datetime
from reports.sections.methodology import add_methodology_section


def link_callback(uri, rel):
    """
    Convert HTML URIs to absolute system paths so xhtml2pdf can access those resources
    """
    sUrl = settings.STATIC_URL      # Typically /static/
    sRoot = settings.STATIC_ROOT    # Typically /home/userX/project_static/
    mUrl = settings.MEDIA_URL       # Typically /media/
    mRoot = settings.MEDIA_ROOT     # Typically /home/userX/project_media/

    if uri.startswith(mUrl):
        path = os.path.join(mRoot, uri.replace(mUrl, ""))
    elif uri.startswith(sUrl):
        path = os.path.join(sRoot, uri.replace(sUrl, ""))
    else:
        return uri

    # make sure that file exists
    if not os.path.isfile(path):
            raise Exception(
                'media URI must start with %s or %s' % (sUrl, mUrl)
            )
    return path

def generate_report_file(project, findings, format='pdf'):
    context = {
        'project': project,
        'findings': findings,
        'today': datetime.now().date(),
    }

    filename = f"{project.project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.{format}"
    output_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if format == 'pdf':
        try:
            html = render_to_string('reports/report_template.html', context)
            with open(output_path, "wb") as out:
                pisa_status = pisa.CreatePDF(html, dest=out, link_callback=link_callback)
                if pisa_status.err:
                    print("Failed to generate PDF:", pisa_status.err)
        except Exception as e:
            print("Error generating PDF:", e)

    elif format == 'docx':
        try:
            doc = Document()

            # Cover Page
            doc.add_heading(f'{project.project_name} - Penetration Testing Report', 0)
            doc.add_paragraph(f"Organization: {project.organization.name}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("Classification: Confidential")
            doc.add_page_break()

            # Description
            doc.add_heading("Description", level=1)
            doc.add_paragraph(project.description or "-")

            # Scope
            doc.add_heading("Scope", level=1)
            doc.add_paragraph(project.scope or "-")
            
            add_methodology_section(doc)

            # Findings Summary
            doc.add_heading("Findings Summary", level=1)
            for finding in findings:
                doc.add_heading(f"{finding.title} ({finding.severity})", level=2)
                doc.add_paragraph(f"Description:\n{finding.description}")
                doc.add_paragraph(f"Impact:\n{finding.impact}")
                doc.add_paragraph(f"Recommendation:\n{finding.recommendation}")
                doc.add_paragraph(f"Status: {finding.status}")
                doc.add_paragraph("-" * 40)

            doc.save(output_path)

        except Exception as e:
            print("Error generating DOCX:", e)

    else:
        raise ValueError("Unsupported format. Only 'pdf' and 'docx' are allowed.")

    print("Report generated at:", output_path)
    return filename
